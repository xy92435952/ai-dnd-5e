"""
模组文件解析服务
支持 PDF / DOCX / Markdown / TXT
"""
import re
from pathlib import Path


async def extract_text(file_path: str, file_type: str) -> str:
    """从上传的文件中提取纯文本"""
    path = Path(file_path)

    if file_type == "pdf":
        return _extract_pdf(path)
    elif file_type == "docx":
        return _extract_docx(path)
    elif file_type in ("md", "markdown"):
        return _extract_markdown(path)
    else:
        return _extract_txt(path)


def _extract_pdf(path: Path) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _extract_markdown(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    # 去掉 markdown 标记，保留纯文本
    text = re.sub(r"#{1,6}\s+", "", raw)           # 标题
    text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)  # 粗体/斜体
    text = re.sub(r"`{1,3}[^`]*`{1,3}", "", text)   # 代码块
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)      # 图片
    text = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", text)  # 链接
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)  # 列表符号
    text = re.sub(r"\n{3,}", "\n\n", text)           # 多余空行
    return text.strip()


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def get_file_type(filename: str) -> str:
    """根据文件名推断文件类型"""
    ext = Path(filename).suffix.lower().lstrip(".")
    mapping = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "md": "md",
        "markdown": "md",
        "txt": "txt",
        "text": "txt",
    }
    return mapping.get(ext, "txt")


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".md", ".markdown", ".txt", ".text"}


def is_allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


def truncate_text(text: str, max_chars: int = 80000) -> str:
    """防止超出 Dify 节点 token 限制"""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[...内容过长，已截断...]"
